import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    run = title.add_run("SkillSphere — Project Structure & Role-Based Overview")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(99, 102, 241) # Indigo color matching design theme
    
    doc.add_paragraph("This document provides a clean, simplified guide to the file structure, user roles, and dashboard views available in the SkillSphere enterprise platform.")
    
    # Section: Folder Tree
    doc.add_heading("📂 Project Directory Layout", level=1)
    
    # ASCII Folder Tree
    tree_text = """SkillSphere/
│
├── config.php                 # Core site and database credentials
├── db.php                     # Database connection helper functions
├── auth.php                   # User session handling & role check
├── helpers.php                # Text truncation and level badges
├── middleware.php             # requireAdmin() / requireInstructor() filters
├── session.php                # Session lifetimes and CSRF fields
│
├── index.php                  # Public Home Page (Trending Courses catalog)
├── courses.php                # Public Catalog (With search bar)
├── course-details.php         # Course details, curriculum list, reviews
│
├── login.php                  # Sign-in page (Role-based redirect)
├── register.php               # Sign-up page (Students & Instructors)
│
├── student/                   # Student Panel (Authenticated Student role)
│   ├── dashboard.php          # Progress metrics & active courses
│   ├── progress.php           # Completed vs total lectures
│   ├── achievements.php       # Unlocked achievement badges
│   ├── leaderboard.php        # Global class rankings
│   └── wishlist.php           # Saved courses
│
├── instructor/                # Instructor Panel (Authenticated Instructor role)
│   ├── dashboard.php          # Earnings overview & course catalog
│   ├── courses.php            # Create new drafts / edit courses
│   ├── students.php           # Directory of registered students
│   ├── analytics.php          # Student sign-up growth chart
│   └── earnings.php           # Gross sales and 70% share metrics
│
├── admin/                     # Admin Panel (Authenticated Admin role)
│   ├── dashboard.php          # System stats & audit logs
│   ├── users/
│   │   └── manage-users.php   # Edit profile / block / ban accounts
│   ├── courses/
│   │   ├── manage-courses.php # Delete / modify course parameters
│   │   └── approvals.php      # Moderate instructor draft course submissions
│   ├── instructors/
│   │   └── payouts.php        # Process payment splits (70%) to instructors
│   └── content/
│       ├── manage-blog.php    # Publish markdown tutorials
│       └── announcements.php  # Send global alerts to student panels
│
├── api/                       # API Endpoints (Dynamic Ajax requests)
│   ├── ai/
│   │   ├── ai-chat.php        # Handles chatbot queries
│   │   └── learning-assistant # Generates AI roadmaps
│   └── courses/
│       └── fetch-courses.php  # Dynamic course query returns
│
└── uploads/                   # Media Upload Storage
    ├── profiles/              # Student/Instructor profile avatars
    └── thumbnails/            # Course thumbnail graphics"""
    
    p = doc.add_paragraph()
    p_run = p.add_run(tree_text)
    p_run.font.name = 'Courier New'
    p_run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.2)
    
    # Section: Roles
    doc.add_heading("👥 Role-Based Access Map", level=1)
    
    doc.add_heading("1. The Student Role (student/)", level=2)
    p1 = doc.add_paragraph("Students are the platform's core consumers.")
    p1.add_run("\n• Entry point: ").bold = True
    p1.add_run("student/dashboard.php")
    p1.add_run("\n• Key Activities:").bold = True
    p1.add_run("\n   - Enrolling in courses and completing video lectures.")
    p1.add_run("\n   - Tracking course percentage metrics on the Study Progress tracker.")
    p1.add_run("\n   - Earning verified badges in the Achievements grid.")
    p1.add_run("\n   - Viewing platform standing on the global Leaderboard.")

    doc.add_heading("2. The Instructor Role (instructor/)", level=2)
    p2 = doc.add_paragraph("Instructors build and manage the educational content.")
    p2.add_run("\n• Entry point: ").bold = True
    p2.add_run("instructor/dashboard.php")
    p2.add_run("\n• Key Activities:").bold = True
    p2.add_run("\n   - Drafting new course submissions (requires Admin approval to publish).")
    p2.add_run("\n   - Editing syllabus titles, prices, descriptions, and prerequisites.")
    p2.add_run("\n   - Auditing active student enrollments.")
    p2.add_run("\n   - Reviewing the Earnings Report (70% revenue share split).")

    doc.add_heading("3. The Administrator Role (admin/)", level=2)
    p3 = doc.add_paragraph("Administrators manage platform moderation and configuration.")
    p3.add_run("\n• Entry point: ").bold = True
    p3.add_run("admin/dashboard.php")
    p3.add_run("\n• Key Activities:").bold = True
    p3.add_run("\n   - Moderating and approving instructor course drafts.")
    p3.add_run("\n   - Modifying user credentials, editing roles, or banning/blocking accounts.")
    p3.add_run("\n   - Tracking system revenue collections and triggering payouts.")
    p3.add_run("\n   - Writing system blogs, setting FAQs, and sending platform announcements.")

    # Section: Localized Telugu Personas & Course Data
    doc.add_heading("🇮🇳 Telugu States Localized Personas & Courses", level=1)
    
    p_loc = doc.add_paragraph("To suit the regional context of Andhra Pradesh and Telangana, the entire user base, simulated logs, and chatbots are localized with authentic Telugu names and roles:")
    
    doc.add_heading("Instructors Profile Table", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Instructor Name'
    hdr_cells[1].text = 'Specialization / Courses'
    hdr_cells[2].text = 'Platform Role & Bio'
    
    instructors = [
        ("Kalyan Ram", "React 19 & Next.js 15, Advanced System Design", "Full-Stack engineering veteran, React core contributor. Ex-Netflix."),
        ("Dr. Sravani Devi", "AI & Deep Learning Bootcamp, Python Data Science", "AI Research Director at MIT. 15+ years of ML pedagogy experience."),
        ("Venkata Srinivas", "Docker, Kubernetes & AWS DevOps", "Principal Cloud Architect. Specializes in cloud-native migrations."),
        ("Harini Reddy", "UI/UX Design Systems for SaaS", "Product Designer & former Design Lead at Stripe. Specialized in UX research."),
        ("Chaitanya Prasad", "Ethical Hacking & Network Pen Testing", "Certified Ethical Hacker (CEH) & corporate infrastructure threat analyst."),
        ("Lakshmi Prasanna", "Go Backend & gRPC, Database Engineering", "Distributed database performance architect & Go core engineer.")
    ]
    
    for name, courses, bio in instructors:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = courses
        row_cells[2].text = bio
        
    doc.add_heading("Representative Student Accounts", level=2)
    p_stud = doc.add_paragraph()
    p_stud.add_run("The database contains 25 realistic student accounts modeled on popular names in Andhra Pradesh & Telangana, including:")
    p_stud.add_run("\n• Ravi Teja Bhupathi").bold = True
    p_stud.add_run(" (ravi.bhupathi@example.com) — Standard student credentials used for testing.")
    p_stud.add_run("\n• Sai Kiran Reddy").bold = True
    p_stud.add_run(" (sai.reddy@example.com) — Standard student credentials used for testing.")
    p_stud.add_run("\n• Murali Krishna Rao").bold = True
    p_stud.add_run(" (murali.rao@example.com)")
    p_stud.add_run("\n• Sravanthi Chowdary").bold = True
    p_stud.add_run(" (sravanthi.chowdary@example.com)")
    p_stud.add_run("\n• Sai Teja Goud").bold = True
    p_stud.add_run(" (sai.goud@example.com)")

    # Save to workspace
    ws_path = "c:/xampp/htdocs/Skillsphere/SkillSphere_Project_Structure.docx"
    doc.save(ws_path)
    print(f"Saved to workspace: {ws_path}")
    
    # Save to user desktop locations
    desktops = [
        "C:/Users/veera/OneDrive/Desktop",
        "C:/Users/veera/Desktop"
    ]
    saved_desktop = False
    for d in desktops:
        if os.path.exists(d):
            dest_path = os.path.join(d, "SkillSphere_Project_Structure.docx")
            try:
                doc.save(dest_path)
                print(f"Saved to desktop: {dest_path}")
                saved_desktop = True
                break
            except Exception as e:
                print(f"Failed to save to {dest_path}: {e}")
                
    if not saved_desktop:
        print("Warning: Could not save to any Desktop directory. File only saved in workspace.")

if __name__ == "__main__":
    main()
